# utils/save_to_word.py
from pathlib import Path
from datetime import datetime
from typing import Optional
import re

from docx import Document


def _add_paragraph_with_markdown(doc: Document, text: str) -> None:
    """Thêm một paragraph, tự parse **bold** sang run.bold = True."""
    pattern = re.compile(r"\*\*(.*?)\*\*")
    p = doc.add_paragraph()
    pos = 0
    for m in pattern.finditer(text):
        # phần trước **bold**
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        # phần **bold**
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    # phần còn lại
    if pos < len(text):
        p.add_run(text[pos:])


def save_to_word(
    content: str,
    score: Optional[float] = None,
    out_dir: str | Path = "outputs",
) -> Path:

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = out_dir / f"affina_post_{ts}.docx"

    doc = Document()

    # Tách đoạn theo 2 dòng trống
    for block in content.split("\n\n"):
        _add_paragraph_with_markdown(doc, block)

    doc.save(file_path)
    return file_path
